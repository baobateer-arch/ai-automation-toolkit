import os
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

EXPORT_DIR = "exports"

# Register a Chinese TrueType font for reportlab PDF rendering.
# On Chinese Windows at least one of these is always available.
FONT_NAME = "Helvetica"  # safe fallback
_font_candidates = [
    ("C:/Windows/Fonts/simhei.ttf", "SimHei"),
    ("C:/Windows/Fonts/msyh.ttc", "MicrosoftYaHei"),
    ("C:/Windows/Fonts/simsun.ttc", "SimSun"),
]
for _path, _name in _font_candidates:
    if os.path.exists(_path):
        try:
            pdfmetrics.registerFont(TTFont(_name, _path))
            FONT_NAME = _name
            break
        except Exception:
            continue


def generate_pdf(report) -> str:
    os.makedirs(EXPORT_DIR, exist_ok=True)
    filepath = os.path.join(EXPORT_DIR, f"report_{report.id}.pdf")

    doc = SimpleDocTemplate(
        filepath, pagesize=A4,
        topMargin=25*mm, bottomMargin=20*mm,
        leftMargin=22*mm, rightMargin=22*mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle", parent=styles["Title"],
        fontName=FONT_NAME, fontSize=22, spaceAfter=6*mm,
        textColor=HexColor("#1e293b"),
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", parent=styles["Normal"],
        fontName=FONT_NAME, fontSize=10,
        textColor=HexColor("#64748b"), spaceAfter=12*mm,
    )
    heading_style = ParagraphStyle(
        "SectionHead", parent=styles["Heading2"],
        fontName=FONT_NAME, fontSize=14,
        spaceBefore=8*mm, spaceAfter=4*mm,
        textColor=HexColor("#2563eb"),
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontName=FONT_NAME, fontSize=11, leading=18,
        spaceAfter=3*mm,
    )
    list_style = ParagraphStyle(
        "ListItem", parent=body_style,
        fontName=FONT_NAME, leftIndent=12, bulletIndent=0,
    )

    content = []

    content.append(Paragraph("AI 文档分析报告", title_style))
    content.append(Paragraph(
        f"文件名：{report.filename} &nbsp;|&nbsp; "
        f"生成时间：{report.created_at.strftime('%Y-%m-%d %H:%M')}",
        subtitle_style,
    ))
    content.append(Spacer(1, 4*mm))

    # Summary
    content.append(Paragraph("摘要", heading_style))
    text = report.summary or "无摘要信息"
    content.append(Paragraph(text, body_style))

    # Key Points
    content.append(Paragraph("关键要点", heading_style))
    points = report.key_points or []
    if points:
        items = [ListItem(Paragraph(f"<bullet>&bull;</bullet>{p}", list_style)) for p in points]
        content.append(ListFlowable(items, bulletType="bullet", leftIndent=16))
    else:
        content.append(Paragraph("暂无", body_style))

    # Risks
    content.append(Paragraph("潜在风险", heading_style))
    risks = report.risks or []
    if risks:
        items = [ListItem(Paragraph(f"<bullet>&bull;</bullet>{r}", list_style)) for r in risks]
        content.append(ListFlowable(items, bulletType="bullet", leftIndent=16))
    else:
        content.append(Paragraph("未识别出明显风险", body_style))

    # Suggestions
    content.append(Paragraph("建议", heading_style))
    suggestions = report.suggestions or []
    if suggestions:
        items = [ListItem(Paragraph(f"<bullet>&bull;</bullet>{s}", list_style)) for s in suggestions]
        content.append(ListFlowable(items, bulletType="bullet", leftIndent=16))
    else:
        content.append(Paragraph("暂无建议", body_style))

    content.append(Spacer(1, 15*mm))
    content.append(Paragraph(
        "由 AI Automation Toolkit 自动生成",
        ParagraphStyle("Footer", parent=styles["Normal"],
                       fontName=FONT_NAME, fontSize=8,
                       textColor=HexColor("#94a3b8"), alignment=1),
    ))

    doc.build(content)
    return filepath


def generate_docx(report) -> str:
    os.makedirs(EXPORT_DIR, exist_ok=True)
    filepath = os.path.join(EXPORT_DIR, f"report_{report.id}.docx")

    doc = Document()

    # Remove default empty paragraph
    if doc.paragraphs:
        p = doc.paragraphs[0]
        p.text = ""

    # Title
    title = doc.add_heading("AI 文档分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"文件名：{report.filename}  |  "
        f"生成时间：{report.created_at.strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()  # spacer

    # Summary
    doc.add_heading("摘要", level=1)
    doc.add_paragraph(report.summary or "无摘要信息")

    # Key Points
    doc.add_heading("关键要点", level=1)
    points = report.key_points or []
    if points:
        for p in points:
            doc.add_paragraph(p, style="List Bullet")
    else:
        doc.add_paragraph("暂无")

    # Risks
    doc.add_heading("潜在风险", level=1)
    risks = report.risks or []
    if risks:
        for r in risks:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph("未识别出明显风险")

    # Suggestions
    doc.add_heading("建议", level=1)
    suggestions = report.suggestions or []
    if suggestions:
        for s in suggestions:
            doc.add_paragraph(s, style="List Bullet")
    else:
        doc.add_paragraph("暂无建议")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("由 AI Automation Toolkit 自动生成")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(filepath)
    return filepath